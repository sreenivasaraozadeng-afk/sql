from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "出海船员管理系统数据库课程设计介绍.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(102, 112, 133)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"


def set_east_asian_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def set_style_font(style, name="Microsoft YaHei", size=11, color=None, bold=None):
    font = style.font
    font.name = name
    font.size = Pt(size)
    if color is not None:
        font.color.rgb = color
    if bold is not None:
        font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_east_asian_font(run)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D0D5DD"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for tag_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn("w:" + tag_name))
        if node is None:
            node = OxmlElement("w:" + tag_name)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        set_cell_text(hdr[i], header, bold=True, color=INK, size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9.5)
    set_table_widths(table, widths)
    set_table_borders(table)
    set_cell_margins(table)
    doc.add_paragraph()
    return table


def add_paragraph(doc, text="", style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_east_asian_font(r1)
        r1.bold = True
        r1.font.color.rgb = INK
        r2 = p.add_run(text[len(bold_lead):])
        set_east_asian_font(r2)
        r2.font.color.rgb = INK
    else:
        run = p.add_run(text)
        set_east_asian_font(run)
        run.font.color.rgb = INK
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_east_asian_font(run)
        run.font.color.rgb = BLUE if level < 3 else DARK_BLUE
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.color.rgb = INK
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.color.rgb = INK
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_table_borders(table, color="C7D7EA")
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_east_asian_font(r)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_east_asian_font(r2)
    r2.font.color.rgb = INK
    doc.add_paragraph()


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11, color=INK)
    normal = styles["Normal"].paragraph_format
    normal.space_after = Pt(6)
    normal.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 22, INK, 0, 8),
        ("Subtitle", 12, MUTED, 0, 16),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=name.startswith("Heading"))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "出海船员管理系统 | 数据库课程设计说明"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_east_asian_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "数据库课程设计答辩材料"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_east_asian_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("数据库课程设计报告")
    set_east_asian_font(run)
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = BLUE

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("出海船员管理系统")
    set_east_asian_font(run)
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = INK

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("船员资源调度与证书风控平台 | 面向教师答辩的详细介绍文档")

    add_table(
        doc,
        ["项目项", "说明"],
        [
            ("课程定位", "数据库课程设计，重点体现关系模型、约束、索引、视图、业务流程和可视化报表。"),
            ("技术路线", "MySQL 8 + FastAPI + SQLAlchemy + 静态 HTML/CSS/JavaScript。"),
            ("系统角色", "管理员、业务经理、证书管理员、船东、船员。"),
            ("核心主线", "船东发布岗位需求，业务经理匹配船员，证书管理员审核证书，系统记录派遣和海历。"),
            ("数据库规模", "16 张业务表、4 个统计视图、演示数据覆盖完整成功/失败流程。"),
        ],
        [1.3, 5.2],
        header_fill=LIGHT_BLUE,
    )

    add_callout(
        doc,
        "给老师的第一句话",
        "本项目不是简单的船员增删改查，而是围绕“船员是否具备上船资格、岗位是否能被匹配、派遣流程是否可追溯”展开的数据库课程设计。我们把船员、证书、岗位、船舶、航线、派遣、海历和日志拆成清晰的关系模型，并通过视图和统计接口把数据库设计结果可视化展示出来。",
    )

    doc.add_page_break()


def build_document():
    doc = Document()
    setup_document(doc)
    title_page(doc)

    add_heading(doc, "一、项目背景与选题意义", 1)
    add_paragraph(
        doc,
        "本课程设计选择“出海船员管理系统”作为题目，是因为船员派遣业务天然具有多实体、多状态、多约束和强审计的特点，非常适合用关系数据库表达。现实场景中，船员能否上船不仅取决于个人档案是否完整，还取决于岗位是否匹配、证书是否审核通过、证书是否过期、船东是否确认派遣、是否已有进行中的出海任务等条件。"
    )
    add_paragraph(
        doc,
        "如果只做一个普通的船员信息管理系统，数据库层面只能体现单表 CRUD，课程设计深度会比较弱。因此本系统进一步扩展为“船员资源调度与证书风控平台”，把业务流程拆成多个相互关联的数据库表，通过外键、唯一约束、检查约束、索引和视图来体现数据库设计能力。"
    )
    add_callout(
        doc,
        "项目亮点",
        "系统的重点不是界面多华丽，而是数据库结构清楚、关系完整、约束合理、查询能支撑业务决策，并能用可视化表格向老师展示数据库设计的结果。",
    )

    add_heading(doc, "二、系统总体目标", 1)
    add_paragraph(doc, "系统设计目标可以概括为四句话：")
    for item in [
        "建立船员、证书、岗位、船舶、航线、派遣和海历的完整数据库模型。",
        "通过证书审核和到期预警控制船员上船风险。",
        "通过岗位匹配评分帮助业务经理选择合适船员。",
        "通过派遣状态日志、操作日志和统计视图体现业务可追溯性和数据库查询能力。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "三、需求分析", 1)
    add_heading(doc, "3.1 用户角色需求", 2)
    add_table(
        doc,
        ["角色", "主要职责", "典型操作"],
        [
            ("管理员", "拥有系统最高权限，负责基础数据维护和整体演示。", "查看统计、维护用户和基础数据、查看日志。"),
            ("业务经理", "负责船员档案、岗位匹配、派遣发起和上/下船确认。", "新增船员、查看匹配分、发起派遣、确认上船/下船。"),
            ("证书管理员", "负责证书录入、审核和风险预警。", "录入证书、审核通过/拒绝、查看到期证书。"),
            ("船东", "发布岗位需求并确认派遣。", "发布岗位、查看自己的派遣、确认派遣。"),
            ("船员", "查看个人档案和出海记录。", "查看个人信息、查看海历。"),
        ],
        [0.95, 2.65, 2.9],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "3.2 核心业务流程", 2)
    for item in [
        "船东维护船舶和岗位需求：岗位需求包含船名、航线、所需岗位、招聘人数、预计上船时间和所需证书。",
        "证书管理员录入并审核证书：证书状态分为待审核、审核通过、审核拒绝；只有审核通过且未过期的证书参与匹配。",
        "业务经理进行智能匹配：系统根据岗位、证书满足度、证书有效期风险、历史海历经验给船员打分。",
        "业务经理发起派遣：派遣状态先进入“待船东确认”。",
        "船东确认派遣：确认后船员进入“待上船”状态，岗位进入“已匹配”状态。",
        "业务经理确认上船：船员进入“出海中”，系统生成海历记录。",
        "业务经理确认下船：船员恢复“在岸可派遣”，海历补全下船时间，岗位关闭。",
    ]:
        add_number(doc, item)

    add_heading(doc, "四、数据库总体设计", 1)
    add_paragraph(
        doc,
        "数据库采用关系模型设计，核心思想是把稳定的基础概念抽象为实体表，把多对多或一对多业务关系拆成关系表，把流程变化写入日志表，把统计展示抽象为视图。这样既能避免数据冗余，又能让业务规则落在清晰的数据库结构中。"
    )
    add_heading(doc, "4.1 表结构分层", 2)
    add_table(
        doc,
        ["分层", "包含表", "设计目的"],
        [
            ("用户与权限层", "users", "统一管理登录账号、角色和显示名称。"),
            ("基础字典层", "positions、certificate_types、ports", "把岗位、证书类型、港口这些可复用数据从业务表中拆出来。"),
            ("船舶航线层", "ship_companies、ships、routes", "描述船东、船舶和航线，为岗位需求提供外键支撑。"),
            ("船员证书层", "crews、certificates、certificate_review_records", "描述船员档案、证书信息和审核过程。"),
            ("岗位派遣层", "job_demands、job_required_certificates、dispatches、dispatch_status_logs", "描述岗位需求、证书要求、派遣流程和状态流转。"),
            ("海历审计层", "voyage_records、operation_logs", "保存上船/下船海历和关键操作日志。"),
        ],
        [1.15, 2.7, 2.65],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "4.2 主要实体关系说明", 2)
    add_paragraph(
        doc,
        "系统中的主要关系包括：一个用户可以对应一个船员档案；一个船东用户可以拥有多个航运公司，一个航运公司拥有多艘船；一个港口可以作为多条航线的出发港或目的港；一个岗位需求关联一个船舶、一条航线、一个岗位，并可要求多个证书；一个派遣记录连接一个岗位需求和一个船员；一个派遣上船后生成一条海历记录。"
    )
    add_paragraph(
        doc,
        "这种设计使每条业务数据都有明确来源。例如岗位需求中的“船名”和“航线”既保留了名称字段方便前端显示，也通过 `ship_id`、`route_id` 关联到规范化实体，便于数据库课程设计中说明外键关系和实体拆分。"
    )

    add_heading(doc, "五、核心数据表介绍", 1)
    add_table(
        doc,
        ["表名", "中文含义", "关键字段", "说明"],
        [
            ("users", "用户表", "username、password_hash、role", "负责登录认证和角色权限控制。"),
            ("crews", "船员表", "user_id、position_id、id_card、status", "保存船员基础档案，身份证唯一。"),
            ("certificates", "证书表", "crew_id、certificate_type_id、expires_at、review_status", "保存证书信息、有效期和审核状态。"),
            ("job_demands", "岗位需求表", "owner_user_id、ship_id、route_id、position_id、status", "船东发布的岗位需求，是匹配和派遣的起点。"),
            ("dispatches", "派遣表", "job_id、crew_id、status", "连接岗位和船员，表示一次派遣流程。"),
            ("voyage_records", "海历表", "dispatch_id、crew_id、onboard_at、offboard_at", "上船后生成，下船后补全。"),
            ("operation_logs", "操作日志表", "user_id、action、target_type、created_at", "记录关键操作，用于审计和展示。"),
        ],
        [1.15, 1.05, 2.25, 2.05],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "六、约束、索引与规范化设计", 1)
    add_heading(doc, "6.1 约束设计", 2)
    for item in [
        "唯一约束：`users.username` 防止账号重复，`crews.id_card` 防止同一身份证重复建档，`certificates.certificate_no` 防止证书编号重复。",
        "外键约束：船员关联用户和岗位，证书关联船员和证书类型，岗位需求关联船舶/航线/岗位，派遣关联岗位和船员。",
        "检查约束：角色、船员状态、证书审核状态、岗位状态、派遣状态都使用 CHECK 约束限制合法值。",
        "级联删除：证书审核记录依赖证书，岗位所需证书依赖岗位需求，适合使用 ON DELETE CASCADE 保持数据一致。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6.2 索引设计", 2)
    add_paragraph(
        doc,
        "系统在高频查询字段上建立索引，例如船员状态 `idx_crews_status`、岗位状态 `idx_job_demands_status`、证书到期日期 `idx_certificates_expires_at`、证书审核状态 `idx_certificates_review_status`、派遣状态 `idx_dispatches_status`、操作日志时间 `idx_operation_logs_created_at`。这些索引直接服务于列表筛选、到期预警、派遣跟踪和统计查询。"
    )

    add_heading(doc, "6.3 规范化说明", 2)
    add_paragraph(
        doc,
        "数据库整体符合第三范式的基本思想：用户、岗位、证书类型、港口、船舶、航线等概念独立成表，避免把大量重复文本直接塞进业务表；岗位和证书之间通过 `job_required_certificates` 表表达一对多关系；派遣状态变化和证书审核过程分别通过日志表记录，避免覆盖历史状态。"
    )
    add_callout(
        doc,
        "为什么仍然保留部分名称字段",
        "在 `job_demands` 和 `voyage_records` 中保留 `ship_name`、`route`、`position` 等名称字段，是为了形成业务快照。即使后续船舶或航线名称发生变化，历史派遣和海历记录仍能保留当时的真实显示信息。这属于有意识的业务冗余，不是随意重复。"
    )

    add_heading(doc, "七、视图与可视化表格", 1)
    add_table(
        doc,
        ["视图名", "展示内容", "前端或报告用途"],
        [
            ("v_crew_certificate_status", "船员、证书、审核状态、到期风险", "证书风控说明和预警表格。"),
            ("v_dispatch_flow_stats", "不同派遣状态的数量", "派遣流程统计。"),
            ("v_route_workload", "每条航线的海历数量和在船数量", "航线工作量排行。"),
            ("v_job_match_overview", "岗位需求、可匹配人数、所需证书数量", "岗位匹配概览。"),
        ],
        [1.9, 2.25, 2.35],
        header_fill=LIGHT_BLUE,
    )
    add_paragraph(
        doc,
        "前端统计首页没有使用复杂图表库，而是用统计卡片、表格和进度条展示数据库查询结果。这样做的好处是演示时更稳定，同时老师能更清楚地看到可视化数据来自哪些表和哪些查询。"
    )

    add_heading(doc, "八、后端接口与业务规则", 1)
    add_heading(doc, "8.1 接口分组", 2)
    add_table(
        doc,
        ["接口组", "代表接口", "作用"],
        [
            ("登录认证", "/api/login、/api/auth/login", "登录并返回 token 和用户信息。"),
            ("船员管理", "/api/crews", "船员列表、新增、编辑、停用。"),
            ("证书管理", "/api/certificates、/api/certificates/{id}/review", "证书录入、审核和预警。"),
            ("基础数据", "/api/ships、/api/ports、/api/routes、/api/positions", "维护船舶、港口、航线和岗位。"),
            ("岗位匹配", "/api/jobs、/api/jobs/{id}/matches", "发布岗位并计算推荐船员。"),
            ("派遣流程", "/api/dispatches/{id}/confirm/onboard/offboard/cancel", "推进派遣状态。"),
            ("统计展示", "/api/dashboard/*", "为统计首页提供数据。"),
            ("日志审计", "/api/operation-logs", "查看关键操作记录。"),
        ],
        [1.15, 2.65, 2.7],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "8.2 匹配评分规则", 2)
    add_paragraph(
        doc,
        "智能匹配不是机器学习，而是规则评分，便于课程设计讲解。总分 100 分，主要由四部分构成：岗位匹配 40 分，证书满足度 40 分，证书有效期风险 10 分，历史海历经验 10 分。系统只返回达到一定分数的船员，业务经理可以看到匹配原因，例如“岗位完全匹配”“所需证书齐全且已审核”“证书有效期充足”“有相近岗位或航线海历”。"
    )
    add_table(
        doc,
        ["评分项", "分值", "数据库依据", "解释"],
        [
            ("岗位匹配", "40", "crews.position 与 job_demands.required_position", "岗位一致才具备基本派遣条件。"),
            ("证书满足度", "40", "certificates 与 job_required_certificates", "要求证书必须存在、审核通过且未过期。"),
            ("有效期风险", "10", "certificates.expires_at", "有效期越充足，风险越低。"),
            ("历史海历", "10", "voyage_records", "有相近岗位或航线经历可加分。"),
        ],
        [1.2, 0.7, 2.25, 2.35],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "九、前端页面与演示流程", 1)
    add_table(
        doc,
        ["页面", "展示内容", "对应数据库重点"],
        [
            ("统计首页", "船员总数、出海人数、证书预警、派遣趋势、航线工作量", "聚合查询、视图、统计接口。"),
            ("船员档案", "船员列表、岗位、状态、新增和停用", "crews、positions、users。"),
            ("证书审核", "证书列表、审核通过/拒绝、到期风险", "certificates、certificate_review_records。"),
            ("船舶航线", "船舶、港口、航线维护", "ship_companies、ships、ports、routes。"),
            ("岗位匹配", "岗位需求、匹配分、发起派遣", "job_demands、job_required_certificates、certificates。"),
            ("派遣跟踪", "确认、上船、下船、取消和状态日志", "dispatches、dispatch_status_logs、voyage_records。"),
            ("操作日志", "关键操作审计", "operation_logs。"),
        ],
        [1.2, 2.8, 2.5],
        header_fill=LIGHT_BLUE,
    )
    add_heading(doc, "9.1 建议课堂演示步骤", 2)
    for item in [
        "使用管理员账号登录，先展示统计首页，让老师看到系统不只是 CRUD。",
        "进入证书审核页面，说明证书审核状态会影响岗位匹配。",
        "进入岗位匹配页面，选择一个岗位，展示匹配分和匹配原因。",
        "发起派遣后进入派遣跟踪页面，依次演示确认、上船、下船。",
        "最后进入操作日志和海历记录页面，说明系统对流程有追溯能力。",
    ]:
        add_number(doc, item)

    add_heading(doc, "十、测试与运行说明", 1)
    add_paragraph(
        doc,
        "项目支持两种运行方式：正式演示可以使用 MySQL 和 Docker Compose；如果本机 MySQL 环境不稳定，也可以使用 SQLite 演示版 `run_sqlite.py`。SQLite 版本会自动按当前模型重建演示数据库，便于课堂演示。"
    )
    add_table(
        doc,
        ["测试类别", "测试内容", "预期结果"],
        [
            ("登录权限", "管理员、经理、证书管理员、船东、船员登录", "返回 token，角色权限正确。"),
            ("证书审核", "待审核证书通过或拒绝", "证书状态改变，生成审核记录。"),
            ("匹配规则", "证书未审核、过期或岗位不符的船员参与匹配", "不推荐或低分，推荐原因清楚。"),
            ("派遣流程", "发起、确认、上船、下船、取消", "状态正确流转，写入状态日志和操作日志。"),
            ("统计页面", "访问 dashboard 接口", "统计卡片和表格数据正常显示。"),
        ],
        [1.2, 2.7, 2.6],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "十一、四人小组分工", 1)
    add_table(
        doc,
        ["成员", "职责", "具体任务", "答辩时可讲内容"],
        [
            ("组长", "数据库总设计与整合", "ER 图、数据字典、init.sql、演示数据、最终整合。", "重点讲数据库结构、约束、视图和设计理由。"),
            ("组员 A", "后端接口与业务规则", "模型、schema、service、路由、匹配评分、派遣日志。", "讲后端如何根据数据库规则完成业务流程。"),
            ("组员 B", "前端页面与可视化", "统计首页、证书审核、岗位匹配、派遣跟踪、表格筛选。", "讲页面如何把数据库查询结果展示给用户。"),
            ("组员 C", "文档、测试与答辩材料", "需求分析、流程图、测试用例、PPT、截图和部署说明。", "讲测试覆盖、演示流程和项目总结。"),
        ],
        [0.85, 1.35, 2.35, 1.95],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "十二、项目总结", 1)
    add_paragraph(
        doc,
        "本系统围绕数据库课程设计的要求，将一个简单船员管理题目扩展成包含多实体、多关系、多状态、多日志和多统计视图的综合数据库应用。系统中的每个功能都能对应到明确的数据表和关系：船员档案体现实体建模，证书审核体现状态控制和历史记录，岗位匹配体现多表查询和业务规则，派遣流程体现状态机和日志追踪，统计首页体现数据库聚合查询和视图设计。"
    )
    add_paragraph(
        doc,
        "如果从老师角度评价，本项目的主要价值在于：数据库不是被动存数据，而是主动支撑业务规则、风险控制和可视化分析。我们通过清晰的表结构、约束、索引、视图和演示流程，把数据库设计能力具体展示出来。"
    )

    add_heading(doc, "附录：答辩时可直接使用的介绍稿", 1)
    add_paragraph(
        doc,
        "老师您好，我们小组的数据库课程设计题目是《出海船员管理系统》。我们没有只做简单的船员增删改查，而是围绕真实的船员派遣业务，把系统设计成船员资源调度与证书风控平台。系统中包括船员、证书、岗位、船舶、港口、航线、派遣、海历和日志等实体。数据库一共设计了 16 张业务表和 4 个统计视图，重点体现关系模型、外键约束、唯一约束、状态约束、索引优化和统计查询。"
    )
    add_paragraph(
        doc,
        "业务主线是：船东发布岗位需求，证书管理员审核船员证书，业务经理根据岗位和证书情况进行智能匹配，船东确认派遣，船员上船后系统生成海历，下船后补全海历记录。整个流程中的状态变化都会写入日志，方便追溯。前端统计首页展示了船员状态、证书预警、派遣趋势和航线工作量，这些可视化数据都来自数据库查询和视图。"
    )
    add_paragraph(
        doc,
        "因此，本项目的重点是用数据库设计支撑业务流程，而不是单纯做页面。我们希望通过这个系统展示对数据库表结构设计、关系建模、数据完整性约束和查询统计的理解。"
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_document()
    print(path)
